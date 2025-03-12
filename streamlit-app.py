import re
import logging
import datetime
import json
import hashlib
import os
import time
import random
from collections import defaultdict
from typing import List, Dict, Tuple, Any, Union
import sqlite3
from jsonschema import validate, ValidationError
from docx import Document
from docx.shared import Inches
import streamlit as st
import google.generativeai as genai
from google.generativeai.types import StopCandidateException
import tornado
import mimetypes


# Konstanten und globale Einstellungen
MODEL_NAME_TEXT = "gemini-2.0-flash-thinking-exp-01-21"  # Für Text
MODEL_NAME_VISION = "gemini-2.0-flash-thinking-exp-01-21"  # Für Bilder

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

API_SLEEP_SECONDS = 60       # Wartezeit zwischen API-Aufrufen
API_MAX_RETRIES = 3          # Maximale Wiederholungsanzahl bei API-Fehlern
SUMMARY_SLEEP_SECONDS = 10  # Pause nach der Zusammenfassung

AUDIT_LOG_FILE = "audit_log.txt"
EXPIRATION_TIME_SECONDS = 300
ROLE_PERMISSIONS = {
    "user": ["REQ", "DATA"],
    "admin": ["REQ", "DATA", "CALC", "IF", "AI"]
}
PRIORITY_MAP = {"HIGH": 1, "MEDIUM": 2, "LOW": 3}

USER_DATA_FILE = "user_data.json"
DISCUSSION_DB_FILE = "discussion_data.db"
RATING_DATA_FILE = "rating_data.json"
AGENT_CONFIG_FILE = "agent_config.json"

USER_DATA_SCHEMA = {
    "type": "object",
    "patternProperties": {
        "^[a-zA-Z0-9_-]+$": {
            "type": "object",
            "properties": {
                "password": {"type": "string"}
            },
            "required": ["password"]
        }
    }
}

AGENT_CONFIG_SCHEMA = {
    "type": "array",
    "items": {
        "type": "object",
        "properties": {
            "name": {"type": "string"},
            "personality": {"type": "string", "enum": ["kritisch", "visionär", "konservativ", "neutral"]},
            "description": {"type": "string"}
        },
        "required": ["name", "personality", "description"]
    }
}


translations = {
    "de": {
        "title": "CipherCore Agenten-Konversation",
        "subtitle": "AI-THINK-TANK – Die KI-Plattform für bahnbrechende Innovationen.",
        "description": "Ein Bild, eine Idee – und in Minuten entsteht eine visionäre Lösung. Von Software über Städteplanung bis hin zu neuen Geschäftsmodellen – du gibst den Impuls, die KI erschafft das Konzept!",
        "login": "Login",
        "register": "Registrierung",
        "username": "Nutzername",
        "password": "Passwort",
        "login_btn": "Login",
        "register_btn": "Registrieren",
        "agent_selection": "Agenten Auswahl",
        "topic": "Diskussionsthema",
        "iterations": "Anzahl Gesprächsrunden",
        "expertise_level": "Experten-Level",
        "language": "Sprache",
        "upload_file": "Datei hochladen (optional)",
        "start_conversation": "Konversation starten",
        "save_discussion": "Diskussion speichern",
        "save_as_word": "Chat als Word speichern",
        "download_word": "Word-Datei herunterladen",
        "upvote": "👍 Upvote",
        "downvote": "👎 Downvote",
        "rating_success": "Bewertung erfolgreich",
        "rating_error": "Fehler bei der Bewertung",
        "save_success": "Diskussion gespeichert",
        "save_error": "Fehler beim Speichern",
        "word_error": "Fehler beim Erstellen der Word-Datei",
        "missing_discussion_id": "Diskussions-ID fehlt. Starten Sie zuerst eine Konversation.",
        "login_warning": "Bitte zuerst einloggen.",
        "api_key_warning": "Bitte geben Sie einen API-Schlüssel ein, um die Anwendung zu nutzen.",
        "connection_error": "Verbindungsfehler. Bitte versuche es später erneut.",
        "unexpected_error": "Ein unerwarteter Fehler ist aufgetreten.",
        "no_agents_selected": "Bitte wähle mindestens einen Agenten aus.",
        "api_key_header": "API-Schlüssel",  # Deutscher Schlüssel
        "saved_discussions" : "Gespeicherte Diskussionen",
        "load_discussions" : "Diskussionen laden",
        "agent_conversation" : "Agenten-Konversation",
        "formatted_output" : "Formatierter Output"
    },
    "en": {
        "title": "CipherCore Agent Conversation",
        "subtitle": "AI-THINK-TANK – The AI platform for groundbreaking innovations.",
        "description": "An image, an idea – and in minutes, a visionary solution is created. From software to urban planning to new business models – you provide the impulse, the AI creates the concept!",
        "login": "Login",
        "register": "Registration",
        "username": "Username",
        "password": "Password",
        "login_btn": "Login",
        "register_btn": "Register",
        "agent_selection": "Agent Selection",
        "topic": "Discussion Topic",
        "iterations": "Number of Conversation Rounds",
        "expertise_level": "Expertise Level",
        "language": "Language",
        "upload_file": "Upload File (optional)",
        "start_conversation": "Start Conversation",
        "save_discussion": "Save Discussion",
        "save_as_word": "Save Chat as Word",
        "download_word": "Download Word File",
        "upvote": "👍 Upvote",
        "downvote": "👎 Downvote",
        "rating_success": "Rating successful",
        "rating_error": "Error in rating",
        "save_success": "Discussion saved",
        "save_error": "Error saving",
        "word_error": "Error creating Word file",
        "missing_discussion_id": "Discussion ID is missing. Start a conversation first.",
        "login_warning": "Please log in first.",
        "api_key_warning": "Please enter an API key to use the application.",
        "connection_error": "Connection error. Please try again later.",
        "unexpected_error": "An unexpected error occurred.",
        "no_agents_selected": "Please select at least one agent.",
        "api_key_header": "API Key",  # Englischer Schlüssel
        "saved_discussions" : "Saved Discussions",
        "load_discussions" : "Load Discussions",
        "agent_conversation" : "Agent Conversation",
        "formatted_output" : "Formatted Output"
    },
    "es": {
        "title": "Conversación de Agentes CipherCore",
        "subtitle": "AI-THINK-TANK – La plataforma de IA para innovaciones revolucionarias.",
        "description": "Una imagen, una idea – y en minutos, se crea una solución visionaria. Desde software hasta planificación urbana y nuevos modelos de negocio – tú das el impulso, la IA crea el concepto!",
        "login": "Iniciar sesión",
        "register": "Registro",
        "username": "Nombre de usuario",
        "password": "Contraseña",
        "login_btn": "Iniciar sesión",
        "register_btn": "Registrarse",
        "agent_selection": "Selección de Agentes",
        "topic": "Tema de Discusión",
        "iterations": "Número de Rondas de Conversación",
        "expertise_level": "Nivel de Experticia",
        "language": "Idioma",
        "upload_file": "Subir Archivo (opcional)",
        "start_conversation": "Iniciar Conversación",
        "save_discussion": "Guardar Discusión",
        "save_as_word": "Guardar Chat como Word",
        "download_word": "Descargar Archivo Word",
        "upvote": "👍 Voto Positivo",
        "downvote": "👎 Voto Negativo",
        "rating_success": "Valoración exitosa",
        "rating_error": "Error en la valoración",
        "save_success": "Discusión guardada",
        "save_error": "Error al guardar",
        "word_error": "Error al crear el archivo Word",
        "missing_discussion_id": "Falta el ID de la discusión. Inicie una conversación primero.",
        "login_warning": "Por favor, inicie sesión primero.",
        "api_key_warning": "Por favor, ingrese una clave API para usar la aplicación.",
        "connection_error": "Error de conexión. Por favor, inténtelo de nuevo más tarde.",
        "unexpected_error": "Ocurrió un error inesperado.",
        "no_agents_selected": "Por favor, seleccione al menos un agente.",
        "api_key_header": "Clave API",  # Spanischer Schlüssel
        "saved_discussions" : "Discusiones Guardadas",
        "load_discussions" : "Cargar Discusiones",
        "agent_conversation" : "Conversación de Agentes",
        "formatted_output" : "Salida Formateada"
    },
    "ru": {
        "title": "Беседа Агентов CipherCore",
        "subtitle": "AI-THINK-TANK – Платформа ИИ для революционных инноваций.",
        "description": "Изображение, идея – и за несколько минут создается визионерское решение. От программного обеспечения до городского планирования и новых бизнес-моделей – вы даете импульс, ИИ создает концепцию!",
        "login": "Вход",
        "register": "Регистрация",
        "username": "Имя пользователя",
        "password": "Пароль",
        "login_btn": "Вход",
        "register_btn": "Зарегистрироваться",
        "agent_selection": "Выбор Агентов",
        "topic": "Тема Обсуждения",
        "iterations": "Количество Раундов Обсуждения",
        "expertise_level": "Уровень Экспертизы",
        "language": "Язык",
        "upload_file": "Загрузить Файл (опционально)",
        "start_conversation": "Начать Беседу",
        "save_discussion": "Сохранить Обсуждение",
        "save_as_word": "Сохранить Чат как Word",
        "download_word": "Скачать Файл Word",
        "upvote": "👍 Плюс",
        "downvote": "👎 Минус",
        "rating_success": "Оценка успешна",
        "rating_error": "Ошибка в оценке",
        "save_success": "Обсуждение сохранено",
        "save_error": "Ошибка при сохранении",
        "word_error": "Ошибка при создании файла Word",
        "missing_discussion_id": "Отсутствует ID обсуждения. Начните беседу сначала.",
        "login_warning": "Пожалуйста, сначала войдите в систему.",
        "api_key_warning": "Пожалуйста, введите ключ API, чтобы использовать приложение.",
        "connection_error": "Ошибка соединения. Пожалуйста, попробуйте позже.",
        "unexpected_error": "Произошла непредвиденная ошибка.",
        "no_agents_selected": "Пожалуйста, выберите хотя бы одного агента.",
        "api_key_header": "Ключ API",  # Russischer Schlüssel
        "saved_discussions" : "Сохраненные Обсуждения",
        "load_discussions" : "Загрузить Обсуждения",
        "agent_conversation" : "Беседа Агентов",
        "formatted_output" : "Форматированный Вывод"
    },
    "zh": {
        "title": "CipherCore 代理对话",
        "subtitle": "AI-THINK-TANK – 创新的人工智能平台。",
        "description": "一张图片，一个想法 – 几分钟内就能创建出一个有远见的解决方案。从软件到城市规划再到新的商业模式 – 您提供灵感，人工智能创建概念！",
        "login": "登录",
        "register": "注册",
        "username": "用户名",
        "password": "密码",
        "login_btn": "登录",
        "register_btn": "注册",
        "agent_selection": "代理选择",
        "topic": "讨论主题",
        "iterations": "对话轮数",
        "expertise_level": "专业水平",
        "language": "语言",
        "upload_file": "上传文件（可选）",
        "start_conversation": "开始对话",
        "save_discussion": "保存讨论",
        "save_as_word": "将聊天保存为Word",
        "download_word": "下载Word文件",
        "upvote": "👍 赞同",
        "downvote": "👎 反对",
        "rating_success": "评分成功",
        "rating_error": "评分错误",
        "save_success": "讨论已保存",
        "save_error": "保存错误",
        "word_error": "创建Word文件时出错",
        "missing_discussion_id": "缺少讨论ID。请先开始对话。",
        "login_warning": "请先登录。",
        "api_key_warning": "请输入API密钥以使用应用程序。",
        "connection_error": "连接错误。请稍后再试。",
        "unexpected_error": "发生意外错误。",
        "no_agents_selected": "请至少选择一个代理。",
        "api_key_header": "API 密钥",  # Chinesischer Schlüssel
        "saved_discussions" : "已保存的讨论",
        "load_discussions" : "加载讨论",
        "agent_conversation" : "代理对话",
        "formatted_output" : "格式化输出"
    }
}

def get_translation(lang: str, key: str) -> str:
    """Holt die Übersetzung für einen Schlüssel, mit speziellem Handling für API-Schlüssel-Warnung."""
    try:
        # Zuerst versuchen, die Übersetzung in der angeforderten Sprache zu holen.
        translation = translations[lang][key]
        if key == "api_key_warning":
            translation += f'  [Get an API key here](https://aistudio.google.com/apikey).'
        return translation
    except KeyError:
        logging.warning(f"Übersetzung für Schlüssel '{key}' in Sprache '{lang}' nicht gefunden.")
        # Fallback auf Englisch, wenn die Übersetzung in der angeforderten Sprache fehlt.
        try:
            translation = translations["en"][key]
            if key == "api_key_warning":
                translation += f'  [Get an API key here](https://aistudio.google.com/apikey).'
            return translation
        except KeyError:
            logging.warning(f"Übersetzung für Schlüssel '{key}' auch in Englisch (en) nicht gefunden.")
            # Erst jetzt, wenn ALLES fehlschlägt, den "Fehlender Schlüssel" Text zurückgeben.
            return f"Fehlender Schlüssel: {key}"



def load_json_data(filename: str, schema: dict = None) -> Dict[str, Any]:
    """
    Lädt JSON-Daten aus einer Datei und validiert sie gegen ein gegebenes Schema.

    Args:
        filename (str): Der Name der JSON-Datei.
        schema (dict, optional): Das Schema zur Validierung der JSON-Daten.

    Returns:
        Dict[str, Any]: Die geladenen JSON-Daten.
    """
    try:
        with open(filename, "r", encoding="utf-8") as f:
            data = json.load(f)
            if schema:
                validate(instance=data, schema=schema)
            return data
    except FileNotFoundError:
        logging.warning(f"Datei '{filename}' nicht gefunden. Starte mit leeren Daten.")
        return {}
    except json.JSONDecodeError as e:
        logging.error(f"Fehler beim Lesen von '{filename}': Ungültiges JSON-Format. Details: {e}")
        return {}
    except ValidationError as e:
        logging.error(f"Datei '{filename}' entspricht nicht dem erwarteten Schema: {e}")
        return {}

def save_json_data(data: Dict[str, Any], filename: str) -> None:
    """
    Speichert JSON-Daten in einer Datei.

    Args:
        data (Dict[str, Any]): Die zu speichernden Daten.
        filename (str): Der Name der Datei, in der die Daten gespeichert werden sollen.
    """
    try:
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)
    except IOError as e:
        logging.error(f"Fehler beim Schreiben in Datei '{filename}': {e}")

def load_user_data() -> Dict[str, Any]:
    """
    Lädt Benutzerdaten aus einer JSON-Datei.

    Returns:
        Dict[str, Any]: Die geladenen Benutzerdaten.
    """
    return load_json_data(USER_DATA_FILE, USER_DATA_SCHEMA)

def save_user_data(user_data: Dict[str, Any]) -> None:
    """
    Speichert Benutzerdaten in einer JSON-Datei.

    Args:
        user_data (Dict[str, Any]): Die zu speichernden Benutzerdaten.
    """
    save_json_data(user_data, USER_DATA_FILE)

def load_rating_data() -> Dict[str, Any]:
    """
    Lädt Bewertungsdaten aus einer JSON-Datei.

    Returns:
        Dict[str, Any]: Die geladenen Bewertungsdaten.
    """
    return load_json_data(RATING_DATA_FILE)

def save_rating_data(rating_data: Dict[str, Any]) -> None:
    """
    Speichert Bewertungsdaten in einer JSON-Datei.

    Args:
        rating_data (Dict[str, Any]): Die zu speichernden Bewertungsdaten.
    """
    save_json_data(rating_data, RATING_DATA_FILE)

def load_agent_config() -> List[Dict[str, str]]:
    """
    Lädt die Agentenkonfiguration aus einer JSON-Datei.

    Returns:
        List[Dict[str, str]]: Die geladene Agentenkonfiguration.
    """
    config = load_json_data(AGENT_CONFIG_FILE, AGENT_CONFIG_SCHEMA)
    if not isinstance(config, list):
        logging.error(f"Agentenkonfiguration in '{AGENT_CONFIG_FILE}' ist ungültig oder leer.")
        return []
    return config

def hash_password(password: str) -> str:
    """
    Hasht ein Passwort mit SHA-256.

    Args:
        password (str): Das zu hashende Passwort.

    Returns:
        str: Der gehashte Passwort-String.
    """
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(password: str, hashed_password: str) -> bool:
    """
    Überprüft, ob ein Passwort mit dem gehashten Passwort übereinstimmt.

    Args:
        password (str): Das zu überprüfende Passwort.
        hashed_password (str): Der gehashte Passwort-String.

    Returns:
        bool: True, wenn das Passwort übereinstimmt, sonst False.
    """
    return hash_password(password) == hashed_password

def register_user(username: str, password: str) -> str:
    """
    Registriert einen neuen Benutzer.

    Args:
        username (str): Der Benutzername des neuen Benutzers.
        password (str): Das Passwort des neuen Benutzers.

    Returns:
        str: Eine Nachricht, die den Status der Registrierung angibt.
    """
    if not re.match(r"^[a-zA-Z0-9_-]+$", username):
        return "Ungültiger Nutzername. Nur Buchstaben, Zahlen, '-', '_' erlaubt."
    if len(password) < 8:
        return "Passwort muss mindestens 8 Zeichen lang sein."
    user_data = load_user_data()
    if username in user_data:
        return "Nutzername bereits vergeben."
    user_data[username] = {"password": hash_password(password)}
    save_user_data(user_data)
    return "Registrierung erfolgreich."

def login_user(username: str, password: str) -> Tuple[str, str]:
    """
    Loggt einen Benutzer ein.

    Args:
        username (str): Der Benutzername des Benutzers.
        password (str): Das Passwort des Benutzers.

    Returns:
        Tuple[str, str]: Eine Nachricht, die den Status des Logins angibt, und den Benutzernamen bei Erfolg.
    """
    user_data = load_user_data()
    if username in user_data and verify_password(password, user_data[username]["password"]):
        return "Login erfolgreich.", username
    return "Login fehlgeschlagen.", None

def create_discussion_table():
    """
    Erstellt die Diskussionstabelle in der SQLite-Datenbank, falls sie nicht existiert.
    """
    conn = sqlite3.connect(DISCUSSION_DB_FILE)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS discussions (
            discussion_id TEXT PRIMARY KEY,
            topic TEXT,
            agents TEXT,
            chat_history TEXT,
            summary TEXT,
            user TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()

create_discussion_table()

def save_discussion_data_db(discussion_id: str, topic: str, agents: List[str], chat_history: List[Dict], summary: str, user: str = None) -> None:
    """
    Speichert Diskussionsdaten in der SQLite-Datenbank.

    Args:
        discussion_id (str): Die ID der Diskussion.
        topic (str): Das Thema der Diskussion.
        agents (List[str]): Die Liste der beteiligten Agenten.
        chat_history (List[Dict]): Der Chatverlauf der Diskussion.
        summary (str): Die Zusammenfassung der Diskussion.
        user (str, optional): Der Benutzer, der die Diskussion gestartet hat.
    """
    conn = sqlite3.connect(DISCUSSION_DB_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("""
            INSERT INTO discussions (discussion_id, topic, agents, chat_history, summary, user)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (discussion_id, topic, json.dumps(agents), json.dumps(chat_history), summary, user))
        conn.commit()
    except sqlite3.Error as e:
        logging.error(f"Datenbankfehler beim Speichern der Diskussion '{discussion_id}': {e}")
        conn.rollback()
    finally:
        conn.close()

def load_discussion_data_db(user: str = None) -> Dict[str, Any]:
    """
    Lädt Diskussionsdaten aus der SQLite-Datenbank.

    Args:
        user (str, optional): Der Benutzer, dessen Diskussionen geladen werden sollen.

    Returns:
        Dict[str, Any]: Die geladenen Diskussionsdaten.
    """
    conn = sqlite3.connect(DISCUSSION_DB_FILE)
    cursor = conn.cursor()
    discussions = {}
    try:
        if user:
            cursor.execute("SELECT * FROM discussions WHERE user = ?", (user,))
        else:
            cursor.execute("SELECT * FROM discussions")
        rows = cursor.fetchall()
        for row in rows:
            disc_id, topic, agents_json, chat_history_json, summary, user_name, timestamp = row
            agents = json.loads(agents_json) if agents_json else []
            chat_history = json.loads(chat_history_json) if chat_history_json else []
            discussions[disc_id] = {
                "topic": topic,
                "agents": agents,
                "chat_history": chat_history,
                "summary": summary,
                "user": user_name,
                "timestamp": timestamp
            }
    except sqlite3.Error as e:
        logging.error(f"Datenbankfehler beim Laden der Diskussionen: {e}")
    finally:
        conn.close()
    return discussions

def evaluate_response(response: str) -> str:
    """
    Bewertet eine Antwort basierend auf bestimmten Schlüsselwörtern.

    Args:
        response (str): Die zu bewertende Antwort.

    Returns:
        str: Die Bewertung der Antwort ("schlechte antwort", "gute antwort", "neutral").
    """
    resp_l = response.lower()
    if "wiederhole mich" in resp_l:
        return "schlechte antwort"
    elif "neue perspektive" in resp_l:
        return "gute antwort"
    else:
        return "neutral"

discussion_ratings = defaultdict(lambda: defaultdict(dict), load_rating_data())

def rate_agent_response(discussion_id: str, iteration: int, agent_name: str, rating_type: str) -> None:
    """
    Bewertet die Antwort eines Agenten.

    Args:
        discussion_id (str): Die ID der Diskussion.
        iteration (int): Die Iteration der Antwort.
        agent_name (str): Der Name des Agenten.
        rating_type (str): Der Typ der Bewertung ("upvote" oder "downvote").
    """
    global discussion_ratings
    if agent_name not in discussion_ratings[discussion_id][iteration]:
        discussion_ratings[discussion_id][iteration][agent_name] = {"upvotes": 0, "downvotes": 0}
    if rating_type == "upvote":
        discussion_ratings[discussion_id][iteration][agent_name]["upvotes"] += 1
    elif rating_type == "downvote":
        discussion_ratings[discussion_id][iteration][agent_name]["downvotes"] += 1
    save_rating_data(discussion_ratings)

def generate_pdf_summary_from_bytes(file_bytes: bytes, api_key: str) -> str:
    """
    Generiert eine Zusammenfassung aus PDF-Bytes.

    Args:
        file_bytes (bytes): Die Bytes der PDF-Datei.
        api_key (str): Der API-Schlüssel für den Gemini-Dienst.

    Returns:
        str: Die generierte Zusammenfassung.
    """
    try:
        mime_type = "application/pdf"
        prompt = "Fasse den Inhalt der PDF zusammen. Achte darauf, dass wichtige Daten nicht verloren gehen!"
        contents = [prompt, genai.types.Part.from_bytes(data=file_bytes, mime_type=mime_type)]
        response = call_gemini_api(contents, api_key, model=MODEL_NAME_TEXT)  # Textmodell für PDF
        return response.get("response", "Start der Konversation.")
    except Exception as e:
        logging.error("Fehler beim Generieren der PDF-Zusammenfassung:", exc_info=e)
        return "Fehler beim Verarbeiten der PDF."

def generate_image_summary_from_bytes(file_bytes: bytes, mime_type: str, api_key: str) -> str:
    """
    Generiert eine Beschreibung/Analyse aus Bild-Bytes.

    Args:
        file_bytes (bytes): Die Bytes der Bilddatei.
        mime_type (str): Der MIME-Typ der Bilddatei.
        api_key (str): Der API-Schlüssel für den Gemini-Dienst.

    Returns:
        str: Die generierte Beschreibung/Analyse.
    """
    try:
        prompt = "Beschreibe den Inhalt des Bildes detailliert. Was sind die Hauptelemente, die Stimmung, und mögliche Interpretationen?"
        contents = [prompt, genai.types.Part.from_bytes(data=file_bytes, mime_type=mime_type)]
        response = call_gemini_api(contents, api_key, model=MODEL_NAME_VISION)  # Vision-Modell für Bilder
        return response.get("response", "Start der Konversation.")
    except Exception as e:
        logging.error("Fehler beim Generieren der Bildbeschreibung:", exc_info=e)
        return "Fehler beim Verarbeiten des Bildes."

def call_gemini_api(contents: list, api_key: str, model: str) -> Dict[str, str]:
    """
    Ruft die Gemini API auf, mit Modell-Auswahl.

    Args:
        contents (list): Die Inhalte, die an die API gesendet werden sollen.
        api_key (str): Der API-Schlüssel für den Gemini-Dienst.
        model (str): Das Modell, das verwendet werden soll.

    Returns:
        Dict[str, str]: Die Antwort von der API.
    """
    client = genai.Client(api_key=api_key)
    retries = 0
    wait_time = 1
    max_wait_time = 60
    while retries < API_MAX_RETRIES:
        try:
            logging.info(f"Sende Anfrage an Gemini ({model}): {str(contents)[:100]}... (Versuch {retries + 1})")
            response = client.models.generate_content(model=model, contents=contents)  # Modell wird übergeben
            if not hasattr(response, "text") or not response.text:
                msg = "Leere Antwort von Gemini API."
                logging.warning(msg)
                return {"response": msg}
            return {"response": response.text}
        except Exception as e:
            err_s = str(e)
            logging.error(f"Gemini API Fehler: {err_s}")
            if "429" in err_s:
                retries += 1
                if retries >= API_MAX_RETRIES:
                    return {"response": f"Fehler: Maximale Anzahl an Versuchen erreicht. API-Kontingent wahrscheinlich erschöpft."}
                wait_time = min(wait_time * 2, max_wait_time)
                jitter = random.uniform(0.5, 1.5)
                actual_wait_time = wait_time * jitter
                logging.warning(f"API-Kontingent erschöpft. Warte {actual_wait_time:.2f} Sekunden...")
                time.sleep(actual_wait_time)
            else:
                return {"response": f"Fehler bei Gemini API Aufruf: {err_s}"}
    return {"response": f"Fehler: Maximale Anzahl an Versuchen erreicht ({API_MAX_RETRIES})."}

def generate_summary(text: str, api_key: str) -> str:
    """
    Generiert eine Textzusammenfassung.

    Args:
        text (str): Der zu zusammenfassende Text.
        api_key (str): Der API-Schlüssel für den Gemini-Dienst.

    Returns:
        str: Die generierte Zusammenfassung.
    """
    prompt = f"Fasse den folgenden Text prägnant zusammen:\n\n{text}"
    result = call_gemini_api([prompt], api_key, model=MODEL_NAME_TEXT)  # Textmodell
    if SUMMARY_SLEEP_SECONDS > 0:
        time.sleep(SUMMARY_SLEEP_SECONDS)
    return result.get("response", "Fehler: Keine Zusammenfassung generiert.")

def process_uploaded_file(uploaded_file, api_key: str) -> str:
    """
    Verarbeitet die hochgeladene Datei und gibt eine Zusammenfassung/Beschreibung zurück.

    Args:
        uploaded_file: Die hochgeladene Datei.
        api_key (str): Der API-Schlüssel für den Gemini-Dienst.

    Returns:
        str: Die generierte Zusammenfassung/Beschreibung.
    """
    if uploaded_file is None:
        return "Start der Konversation."

    try:
        file_bytes = uploaded_file.read()
        uploaded_file.seek(0)  # Wichtig: Zurücksetzen des Dateizeigers
        mime_type = uploaded_file.type  # MIME-Type direkt aus Streamlit-Objekt

        if mime_type == "application/pdf":
            return generate_pdf_summary_from_bytes(file_bytes, api_key)
        elif mime_type.startswith("image"):
            return generate_image_summary_from_bytes(file_bytes, mime_type, api_key)
        else:
            logging.warning(f"Nicht unterstützter Dateityp: {mime_type}")
            return f"Nicht unterstützter Dateityp: {mime_type}"

    except Exception as e:
        logging.error(f"Fehler beim Verarbeiten der Datei: {e}", exc_info=e)
        return "Fehler beim Verarbeiten der Datei."

def joint_conversation_with_selected_agents(
    conversation_topic: str,
    selected_agents: List[Dict[str, str]],
    iterations: int,
    expertise_level: str,
    language: str,
    chat_history: List[Dict[str, str]],
    user_state: str,
    discussion_id: str = None,
    api_key: str = None,
    uploaded_file=None
):
    """
    Führt eine Konversation mit ausgewählten Agenten durch.

    Args:
        conversation_topic (str): Das Thema der Konversation.
        selected_agents (List[Dict[str, str]]): Die Liste der ausgewählten Agenten.
        iterations (int): Die Anzahl der Gesprächsrunden.
        expertise_level (str): Das Experten-Level der Agenten.
        language (str): Die Sprache der Konversation.
        chat_history (List[Dict[str, str]]): Der Chatverlauf.
        user_state (str): Der Zustand des Benutzers.
        discussion_id (str, optional): Die ID der Diskussion.  Wird automatisch generiert, wenn nicht angegeben.
        api_key (str, optional): Der API-Schlüssel für den Gemini-Dienst.
        uploaded_file: Die hochgeladene Datei (optional).

    Yields:
        Tuple: Aktualisierter Chatverlauf, formatierter Textabschnitt, Diskussions-ID, Iterationsnummer, Agentenname.
    """
    if discussion_id is None:
        discussion_id = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    # ---  Verwende IMMER eine feste Datei.  Guter Name, um Konflikte zu vermeiden. ---
    chat_history_filename = "chat_history.txt"

    # --- Option zum Auskommentieren (für lokale Entwicklung nützlich) ---
    # Setze USE_CHAT_HISTORY_FILE auf False, um das Schreiben zu deaktivieren.
    USE_CHAT_HISTORY_FILE = False   #  Oder True

    active_agents_names = [sa["name"] for sa in selected_agents]
    num_agents = len(active_agents_names)
    agent_outputs = [""] * num_agents
    topic_changed = False
    logging.info(f"Konversation gestartet: {active_agents_names}, iters={iterations}, level={expertise_level}, lang={language}, Diskussions-ID: {discussion_id}, Datei: {uploaded_file is not None}")

    initial_summary = process_uploaded_file(uploaded_file, api_key)
    current_summary = initial_summary

    for i in range(iterations):
        agent_idx = i % num_agents
        current_agent_name = active_agents_names[agent_idx]
        current_agent_config = next((a for a in selected_agents if a["name"] == current_agent_name), None)
        current_personality = current_agent_config.get("personality", "neutral")
        current_instruction = current_agent_config.get("instruction", "")

        prompt_text = (
            f"Wir führen eine Konversation über: '{conversation_topic}'.\n"
            + (f"Zusätzliche Informationen sind in der angehängten Datei verfügbar.\n" if uploaded_file is not None else "")
            + f"Hier ist die Zusammenfassung der bisherigen Diskussion:\n{current_summary}\n\n"
            + f"Iteration {i+1}: Agent {current_agent_name}, bitte antworte. {current_instruction}\n"
        )
        if i > 0:
            prompt_text += f"Der vorherige Agent sagte: {agent_outputs[(agent_idx - 1) % num_agents]}\n"
        if current_personality == "kritisch":
            prompt_text += "\nSei kritisch."
        elif current_personality == "visionär":
            prompt_text += "\nSei visionär."
        elif current_personality == "konservativ":
            prompt_text += "\nSei konservativ."
        prompt_text += f"\n\nAntworte auf {language}."

        # Inhalte vorbereiten: Datei-Part anhängen, falls vorhanden.
        contents = [prompt_text]
        if uploaded_file is not None:
            try:
                file_bytes = uploaded_file.read()  # Bytes erneut lesen, falls benötigt
                uploaded_file.seek(0)
                mime_type = uploaded_file.type
                contents = [prompt_text, genai.types.Part.from_bytes(data=file_bytes, mime_type=mime_type)]
            except Exception as e:
                logging.error(f"Fehler beim Lesen der Datei für Iteration {i+1}: {e}", exc_info=e)
                yield chat_history, f"Fehler beim Lesen der Datei in Iteration {i+1}.", discussion_id, (i + 1), current_agent_name  # Fehler melden
                continue  # Nächste Iteration

        api_resp = call_gemini_api(contents, api_key, model=MODEL_NAME_TEXT)  # Immer Textmodell in der Konversation
        agent_output = api_resp.get("response", f"Keine Antwort von {current_agent_name}")
        agent_outputs[agent_idx] = agent_output

        chat_history.append({
            "role": "user",
            "content": f"Agent {current_agent_name} (Iteration {i + 1}): Thema {conversation_topic}, Zusammenfassung bis Runde {i}: {current_summary}, Datei: {'vorhanden' if uploaded_file is not None else 'nicht vorhanden'}"
        })
        chat_history.append({
            "role": "assistant",
            "content": f"Antwort von Agent {current_agent_name} (Iteration {i+1}):\n{agent_output}"
        })

        # ---  Nur schreiben, wenn USE_CHAT_HISTORY_FILE True ist  ---
        if USE_CHAT_HISTORY_FILE:
            try:
                with open(chat_history_filename, "w", encoding="utf-8") as f:
                    for message in chat_history:
                        f.write(f"{message['role']}: {message['content']}\n")
            except IOError as e:
                logging.error(f"Fehler beim Schreiben in Chatverlauf-Datei '{chat_history_filename}': {e}")


        new_summary_input = f"Bisherige Zusammenfassung:\n{current_summary}\n\nNeue Antwort von {current_agent_name}:\n{agent_output}"
        current_summary = generate_summary(new_summary_input, api_key)
        time.sleep(API_SLEEP_SECONDS)

        qual = evaluate_response(agent_output)
        if qual == "schlechte antwort":
            logging.info(f"{current_agent_name} => 'schlechte antwort', retry ...")
            # Retry-Logik mit korrekter Behandlung der Datei
            retry_contents = ["Versuche eine kreativere Antwort."]
            if uploaded_file is not None:
                try:
                    file_bytes = uploaded_file.read()  # Bytes für Retry erneut lesen
                    uploaded_file.seek(0)
                    mime_type = uploaded_file.type
                    retry_contents = [genai.types.Part.from_bytes(data=file_bytes, mime_type=mime_type), "Versuche eine kreativere Antwort."]
                except Exception as e:
                    logging.error(f"Fehler beim Lesen der Datei für Retry: {e}", exc_info=e)
                    yield chat_history, "Fehler beim Lesen der Datei während des Retrys.", discussion_id, (i + 1), current_agent_name
                    continue
            retry_resp = call_gemini_api(retry_contents, api_key, model=MODEL_NAME_TEXT)  # Textmodell
            retry_output = retry_resp.get("response", f"Keine Retry-Antwort von {current_agent_name}")
            if "Fehler bei Gemini API Aufruf" not in retry_output:
                agent_output = retry_output
            agent_outputs[agent_idx] = agent_output

        # WICHTIG: rating_info VOR dem Yield aktualisieren
        st.session_state['rating_info']["discussion_id"] = discussion_id
        st.session_state['rating_info']["iteration"] = i + 1  # Korrekte Iterationsnummer
        st.session_state['rating_info']["agent_name"] = current_agent_name

        logging.info(f"Antwort Agent {current_agent_name} (i={i+1}): {agent_output[:50]}...")
        formatted_output_chunk = (
            f"**Iteration {i+1}: Agent {current_agent_name} ({current_personality})**\n\n"
            f"{agent_output}\n\n"
            "---\n\n"
        )
        yield chat_history, formatted_output_chunk, discussion_id, (i + 1), current_agent_name

        if i > iterations * 0.6 and agent_output == agent_outputs[(agent_idx - 1) % num_agents] and not topic_changed:
            new_topic = "Neues Thema: KI-Trends 2026"
            contents = [new_topic]
            # Neue Themenlogik mit korrekter Behandlung der Datei
            if uploaded_file is not None:
                try:
                    file_bytes = uploaded_file.read()  # Bytes erneut lesen
                    uploaded_file.seek(0)
                    mime_type = uploaded_file.type
                    contents = [new_topic, genai.types.Part.from_bytes(data=file_bytes, mime_type=mime_type)]
                except Exception as e:
                    logging.error(f"Fehler beim Lesen der Datei für neues Thema: {e}", exc_info=e)
                    yield chat_history, "Fehler beim Lesen der Datei beim Themenwechsel.", discussion_id, (i + 1), current_agent_name
                    continue

            agent_outputs = [new_topic] * num_agents
            topic_changed = True

    final_summary_input = "Gesamter Chatverlauf:\n" + "\n".join(
        [f"{m['role']}: {m['content']}" for m in chat_history]
    )
    final_summary = generate_summary(final_summary_input, api_key)
    chat_history.append({
        "role": "assistant",
        "content": f"**Gesamtzusammenfassung**:\n{final_summary}"
    })

    if user_state:
        save_discussion_data_db(discussion_id, conversation_topic, active_agents_names, chat_history, final_summary, user_state)
        logging.info(f"Diskussion {discussion_id} für {user_state} in Datenbank gespeichert.")
    else:
        logging.info("Keine Speicherung in Datenbank, kein Benutzer eingeloggt.")

    final_text = agent_outputs[-1]
    chat_history.append({
        "role": "assistant",
        "content": f"Finale Aussage:\n{final_text}"
    })
    logging.info(f"Finale Aussage: {final_text}")
    yield chat_history, final_summary, discussion_id, None, None

def save_chat_as_word(chat_history: List[Dict], discussion_id: str) -> str:
    """
    Speichert den Chatverlauf als Word-Dokument.

    Args:
        chat_history (List[Dict]): Der Chatverlauf.
        discussion_id (str): Die ID der Diskussion.

    Returns:
        str: Der Dateiname des gespeicherten Word-Dokuments.
    """
    document = Document()
    document.add_heading(f'CipherCore Agenten-Diskussion: {discussion_id}', level=1)
    for message in chat_history:
        role = message['role']
        content = message['content']
        if role == 'user':
            document.add_paragraph("Nutzer:", style='List Bullet').add_run(f" {content}").bold = True
        elif role == 'assistant':
            agent_name_match = re.search(r'Agent (.*?)\s', content)
            agent_name = agent_name_match.group(1) if agent_name_match else "Agent"
            p = document.add_paragraph(f"{agent_name}:", style='List Bullet')
            p.add_run(f" {content.split(':\n', 1)[1] if ':\n' in content else content}")
    filename = f"CipherCore_Diskussion_{discussion_id}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        document.save(filename)
        logging.info(f"Word-Datei '{filename}' erfolgreich gespeichert.")
        return filename
    except Exception as e:
        logging.error(f"Fehler beim Speichern der Word-Datei: {e}")
        return None



def main():
    """
    Hauptfunktion der Streamlit-Anwendung.
    """
    # --- SPRACHAUSWAHL ---
    # 1. Sprache im Session State speichern (mit 'en' als Standard)
    if 'language' not in st.session_state:
        st.session_state['language'] = 'en'

    # 2. Sprachauswahl (Radio-Buttons).  WICHTIG: key='language'
    language_options = {
        "de": "Deutsch",
        "en": "Englisch",
        "fr": "Französisch",
        "es": "Spanisch",
        "ru": "Russisch",
        "zh": "Chinesisch"
    }
    selected_lang_code = st.sidebar.radio(
        "Language",
        options=list(language_options.keys()),
        format_func=lambda code: language_options[code],
        key='language',
        horizontal=True,
        index=list(language_options.keys()).index(st.session_state['language'])  # Standard aus Session State
    )

    lang = selected_lang_code

    # --- Initialisiere ALLE Session State Variablen *vor* dem ersten Aufruf von get_translation ---
    if 'user_state' not in st.session_state:
        st.session_state['user_state'] = None
    if 'chat_history' not in st.session_state:
        st.session_state['chat_history'] = []
    if 'discussion_id' not in st.session_state:
        st.session_state['discussion_id'] = None
    if 'rating_info' not in st.session_state:
        st.session_state['rating_info'] = {}
    if 'formatted_output_text' not in st.session_state:
        st.session_state['formatted_output_text'] = ""
    if 'api_key' not in st.session_state:
        st.session_state['api_key'] = None  # Wichtig: Auch api_key initialisieren!
    if 'uploaded_file' not in st.session_state:
        st.session_state['uploaded_file'] = None


    # --- JETZT erst die Texte holen, nachdem die Sprache korrekt gesetzt wurde ---
    st.title(get_translation(lang, "title"))
    st.markdown(get_translation(lang, "subtitle"))
    st.markdown(get_translation(lang, "description"))


    st.sidebar.header(get_translation(lang, "api_key_header"))  # <--- Hier war das Problem
    api_key = st.sidebar.text_input("Geben Sie Ihren Gemini API-Schlüssel ein:", type="password")
    if not api_key:
        st.warning(get_translation(lang, "api_key_warning"))
        return
    else:
        st.session_state['api_key'] = api_key  # api_key in den Session State, falls eingegeben

    # ... (Der Rest der main-Funktion bleibt gleich, aber stelle sicher, dass du ÜBERALL get_translation verwendest) ...
    with st.expander(get_translation(lang, "login_register"), expanded=False):  # Übersetzter Expander-Text
        col1, col2 = st.columns(2)
        with col1:
            st.subheader(get_translation(lang, "login"))  # Überschriften übersetzen
            username_login = st.text_input(get_translation(lang, "username"), key="username_login")
            password_login = st.text_input(get_translation(lang, "password"), type="password", key="password_login")
            login_btn = st.button(get_translation(lang, "login_btn"))
            login_status = st.empty()
            if login_btn:
                msg, logged_in_user = login_user(username_login, password_login)
                if logged_in_user:
                    st.session_state['user_state'] = logged_in_user
                    login_status.success(msg)  # KEINE Übersetzung der Login-Meldung, da diese dynamisch generiert wird!
                    st.rerun()
                else:
                    login_status.error(msg)  # KEINE Übersetzung

        with col2:
            st.subheader(get_translation(lang, "register"))
            username_register = st.text_input(get_translation(lang, "username"), key="username_register")
            password_register = st.text_input(get_translation(lang, "password"), type="password", key="password_register")
            register_btn = st.button(get_translation(lang, "register_btn"))
            register_status = st.empty()
            if register_btn:
                msg = register_user(username_register, password_register)
                register_status.info(msg)  # KEINE Übersetzung der Registrierungs-Meldung

    st.markdown("---")
    agent_config_data = load_agent_config()
    agent_selections = {}
    st.subheader(get_translation(lang, "agent_selection"))
    with st.expander(get_translation(lang, "agent_selection"), expanded=False):
        for agent_data in agent_config_data:
            agent_selections[agent_data["name"]] = {
                "selected": st.checkbox(agent_data["name"]),  # KEINE Übersetzung des Agenten-Namens
                "personality": st.radio(
                    f"Persönlichkeit für {agent_data['name']}", # Keine Übersetzung innerhalb des f-strings
                    ["kritisch", "visionär", "konservativ", "neutral"], # Keine Übersetzung von festen Werten
                    horizontal=True,
                    key=f"personality_{agent_data['name']}"
                )
            }

    topic_input = st.text_input(get_translation(lang, "topic"))
    iteration_slider = st.slider(get_translation(lang, "iterations"), 1, 50, value=10, step=1)
    level_radio = st.radio(get_translation(lang, "expertise_level"), ["Beginner", "Fortgeschritten", "Experte"], horizontal=True) # Keine Übersetzung der Level
    # --- SPRACHAUSWAHL OBEN, nicht hier ---
    # lang_radio = st.radio("Sprache", ["Deutsch", "Englisch", "Französisch", "Spanisch", "Russisch", "Chinesisch"], horizontal=True)

    st.subheader(get_translation(lang, "upload_file"))
    uploaded_file = st.file_uploader(get_translation(lang, "upload_file"), type=["pdf", "png", "jpg", "jpeg", "gif"]) # Hier wird der übersetzte Text verwendet

    if uploaded_file is not None:
        st.session_state['uploaded_file'] = uploaded_file
        file_type = uploaded_file.type
        st.write(f"Hochgeladener Dateityp: {file_type}")  # Dateityp-Anzeige nicht übersetzen

        if file_type.startswith("image"):
            st.image(uploaded_file)
        elif file_type == "application/pdf":
            st.write("PDF-Datei hochgeladen (Vorschau nicht unterstützt)")  # ODER übersetzen, wenn du möchtest
        else:
            st.write("Datei hochgeladen") # ODER übersetzen.

    with st.expander(get_translation(lang,"saved_discussions"), expanded=False):
        load_disc_btn = st.button(get_translation(lang,"load_discussions"))
        saved_discussions = st.empty()
        if load_disc_btn:
            if st.session_state['user_state']:
                disc_data = load_discussion_data_db(st.session_state['user_state'])
                saved_discussions.json(disc_data) # Keine Übersetzung, da es sich um Daten handelt.
            else:
                saved_discussions.warning(get_translation(lang,"login_warning"))

    start_btn = st.button(get_translation(lang, "start_conversation"))

    st.subheader(get_translation(lang, "agent_conversation"))
    for message in st.session_state['chat_history']:
        with st.chat_message(message["role"]):
            st.markdown(message["content"]) # Keine Übersetzung, da dynamischer Inhalt

    st.subheader(get_translation(lang, "formatted_output"))
    st.markdown(st.session_state['formatted_output_text'])# Keine Übersetzung, da dynamischer Inhalt

    rating_col1, rating_col2, rating_col3 = st.columns([1, 1, 3])
    if st.session_state['rating_info'].get("iteration") is not None:
        with rating_col1:
            upvote_btn = st.button(get_translation(lang, "upvote"))
        with rating_col2:
            downvote_btn = st.button(get_translation(lang, "downvote"))
        with rating_col3:
            rating_label = st.empty()

        if upvote_btn:
            did = st.session_state['rating_info'].get("discussion_id")
            itn = st.session_state['rating_info'].get("iteration")
            agn = st.session_state['rating_info'].get("agent_name")
            if did and itn and agn:
                rate_agent_response(did, itn, agn, "upvote")
                rating_label.success(get_translation(lang, "rating_success"))
            else:
                rating_label.error(get_translation(lang, "rating_error"))
        if downvote_btn:
            did = st.session_state['rating_info'].get("discussion_id")
            itn = st.session_state['rating_info'].get("iteration")
            agn = st.session_state['rating_info'].get("agent_name")
            if did and itn and agn:
                rate_agent_response(did, itn, agn, "downvote")
                rating_label.success(get_translation(lang, "rating_success"))
            else:
                rating_label.error(get_translation(lang, "rating_error"))

    save_col1, save_col2 = st.columns(2)
    with save_col1:
        save_btn = st.button(get_translation(lang, "save_discussion"))
        save_status = st.empty()
        if save_btn:
            if st.session_state['user_state']:
                active_agents_names = [agent['name'] for agent in agent_config_data if agent_selections[agent['name']]['selected']]
                save_discussion_data_db(st.session_state['discussion_id'], topic_input, active_agents_names, st.session_state['chat_history'], "Manuell gespeichert", st.session_state['user_state'])
                save_status.success(get_translation(lang, "save_success"))
            else:
                save_status.warning(get_translation(lang, "login_warning"))
    with save_col2:
        word_save_btn = st.button(get_translation(lang, "save_as_word"))
        if word_save_btn:
            if st.session_state['discussion_id']:
                word_filename = save_chat_as_word(st.session_state['chat_history'], st.session_state['discussion_id'])
                if word_filename:
                    with open(word_filename, "rb") as file:
                        st.download_button(
                            label=get_translation(lang, "download_word"), # Übersetztes Label
                            data=file,
                            file_name=word_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.error(get_translation(lang, "word_error"))
            else:
                st.warning(get_translation(lang, "missing_discussion_id"))

    if start_btn:
        selected_agents = [
            {"name": agent, "personality": agent_selections[agent]["personality"], "instruction": next((a["description"] for a in agent_config_data if a["name"] == agent), "")}
            for agent in agent_selections if agent_selections[agent]["selected"]
        ]
        if not selected_agents:
            st.warning(get_translation(lang, "no_agents_selected")) # Übersetzte Warnung
        else:
            st.session_state['chat_history'] = []
            st.session_state['formatted_output_text'] = ""
            st.session_state['discussion_id'] = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            st.session_state['rating_info'] = {}
            try:
                with st.spinner("Konversation wird gestartet..."):  # Spinner-Text NICHT übersetzen ,aber eventuell auch hierfür einen key erstellen
                    agent_convo = joint_conversation_with_selected_agents(
                        conversation_topic=topic_input,  # Kein Prompt mehr hier
                        selected_agents=selected_agents,
                        iterations=iteration_slider,
                        expertise_level=level_radio, # Keine Übersetzung
                        language=lang,  # Verwende den Sprachcode!
                        chat_history=[],
                        user_state=st.session_state['user_state'],
                        discussion_id=st.session_state['discussion_id'],
                        api_key=st.session_state['api_key'],
                        uploaded_file=st.session_state.get('uploaded_file')
                    )
                    for updated_hist, chunk_text, disc_id, iteration_num, agent_n in agent_convo:
                        st.session_state['discussion_id'] = disc_id
                        st.session_state['rating_info']["discussion_id"] = disc_id
                        st.session_state['rating_info']["iteration"] = iteration_num
                        st.session_state['rating_info']["agent_name"] = agent_n

                        if updated_hist and len(updated_hist) > len(st.session_state['chat_history']):
                            new_messages = updated_hist[len(st.session_state['chat_history']):]
                            for message in new_messages:
                                st.session_state['chat_history'].append(message)
                                with st.chat_message(message["role"]):
                                    st.markdown(message["content"])
                            st.session_state['formatted_output_text'] += chunk_text

            except (tornado.websocket.WebSocketClosedError, tornado.iostream.StreamClosedError) as e:
                st.error(get_translation(lang, "connection_error"))
            except StopCandidateException as e:
                st.error(f"Die Konversation wurde unerwartet beendet: {e}")  # Keine Übersetzung, da es eine spezifische Fehlermeldung ist.
            except Exception as e:
                st.error(get_translation(lang, "unexpected_error"))
if __name__ == "__main__":
    main()
